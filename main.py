from random import choice
import sys
from tkinter import *
from tkinter import scrolledtext
from tkinter import filedialog as fd
import docx
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

root = Tk()
root.title('Ćwiczenia na kodowanie DNA')
root.resizable(False, False)
root.geometry('600x400')
root.iconbitmap('logo.ico')

def clear(event):
    errorLabel['text'] = ''
    wordsEntry.delete("0.0", END)
def assign(l):
    global names, letters, keysToAdd, alphc, keysToAdd, amins, words
    nameChosen = choice(keysToAdd)
    while nameChosen[1]=='stop' or 'start' in nameChosen[1]: nameChosen = choice(keysToAdd)

    letters[nameChosen[1]] = l
    alphc.remove(l)
    keysToAdd.remove(nameChosen)
    for item in names.items():
        if item[1]==nameChosen[1]:
            try:
                keysToAdd.remove(item)
            except: pass
def DNAtoRNA(l):
    vars = {
    'A':'T',
    'U':'A',
    'G':'C',
    'C':'G'
    }
    return vars.get(l)
def translateEntry(*args):
    global names, letters, keysToAdd, alphc, keysToAdd, amins, words, results
    errorLabel['text'] = ''

    results = []
    names = {
        'UUU':'phe','UUC':'phe','UUA':'leu',
        'UUG':'leu','CUU':'leu','CUC':'leu',
        'CUA':'leu','CUG':'leu','AUU':'ile',
        'AUC':'ile','AUA':'ile','AUG':'met (start)',
        'GUU':'val','GUC':'val','GUA':'val',
        'GUG':'val','UCU':'ser','UCC':'ser',
        'UCA':'ser','UCG':'ser','CCU':'pro',
        'CCC':'pro','CCA':'pro','CCG':'pro',
        'ACU':'thr','ACC':'thr','ACA':'thr',
        'ACG':'thr','GCU':'ala','GCC':'ala',
        'GCA':'ala','GCG':'ala','UAU':'tyr',
        'UAC':'tyr','UAA':'stop','UAG':'stop',
        'CAU':'his','CAC':'his','CAA':'gln',
        'CAG':'gln','AAU':'asn','AAC':'asn',
        'AAA':'lys','AAG':'lys','GAU':'asp',
        'GAC':'asp','GAA':'glu','GAG':'glu',
        'UGU':'cys','UGC':'cys','UGA':'stop',
        'UGG':'trp','CGU':'arg','CGC':'arg',
        'CGA':'arg','CGG':'arg','AGU':'ser',
        'AGC':'ser','AGA':'arg','AGG':'arg',
        'GGU':'gly','GGC':'gly','GGA':'gly','GGG':'gly'
    }
    alphc = list('AĄBCĆDEĘFGHIJKLŁMNOÓPQRSŚTUVWXYZŹŻ')
    letters, amins = {}, {}

    words = wordsEntry.get("0.0", END).upper().split('\n')
    try:
        while words[-1]=='': words.pop(-1)
    except:
        errorLabel['text'] = 'Wpisz frazy do zakodowania'; return

    used = []
    for word in words:
        for l in set(word):
            if l not in used:
                used.append(l)
    if len(used) > 20:
        errorLabel['text'] = 'Przekroczono limit znaków'
        return

    keysToAdd = list(names.items())
    for l in used: assign(l)
    for i in range(19-len(used)):
        l = choice(alphc)
        assign(l)
    for l in letters.values(): amins[l] = []

    wordsEntry.insert(END, '\n\nTabela kodowania\n')
    for item in names.items():
        if 'start' not in item[1] and item[1]!='stop':
            amins[letters[item[1]]].append(item[0])
            wordsEntry.insert(END, f'{item[0]} {item[1]} ({letters[item[1]]})\n')
        else:
            wordsEntry.insert(END, f'{item[0]} {item[1]}\n')

    wordsEntry.insert(END, '\n\nZakodowane hasła:\n')
    for i, word in enumerate(words):
        #rna cipher
        rna = ''
        for _ in range(int(frontSpin.get())): rna += choice(['U', 'C', 'A', 'G'])
        rna += 'AUG'
        for l in word: rna += choice(amins[l])
        rna += choice(['UAA', 'UAG', 'UGA'])
        for _ in range(int(backSpin.get())): rna += choice(['U', 'C', 'A', 'G'])

        #dna cipher
        dna = ''
        for l in rna:
            dna += DNAtoRNA(l)
        results.append(dna)
        wordsEntry.insert(END, f'{i+1}. {dna}\n')
def generateFiles(*args):
    global names, letters, results
    errorLabel['text'] = ''

    try:
        file = fd.asksaveasfilename(filetypes=(('Dokument programu MS Word','*.docx'),('Wszystkie pliki', '*.* ')))
        if not file:
            errorLabel['text'] = 'Nie wybrano pliku'; return

        doc = docx.Document()
        section = doc.sections[-1]
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        p = doc.add_paragraph()
        run = p.add_run('Karty pracy do części 1')
        run.bold = True
        run.font.size = Pt(15)

        run = p.add_run('\nTabela kodonów')
        run.bold = True
        run.font.size = Pt(12)

        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        table.style.font.size = Pt(10)
        count,i,j=0,0,0
        for item in names.items():
            if count==4:
                count = 0
                j += 1
            if j==4:
                j=0; i+=1
            if 'start' not in item[1] and item[1]!='stop':
                table.rows[j].cells[i].text +=  f'{item[0]} {item[1]} ({letters[item[1]]})\n'
            else:
                table.rows[j].cells[i].text += f'{item[0]} {item[1]}\n'
            count += 1

        p = doc.add_paragraph()
        run = p.add_run('\nNa podstawie podanej sekwencji DNA przeprowadź transkrypcję i translację.\
 Wyniki zapisz poniżej. Pamiętaj, że zaczynasz od kodonu START, a kończysz na kodnie STOP.')
        run.font.size = Pt(10)

        for r in results:
            run = p.add_run(f'\n\nDNA: {r}\nRNA: {3*len(r)*"."}\nKOD: {3*len(r)*"."}')
            run.font.size = Pt(10)

        doc.save(file)
        errorLabel['text'] = 'Karta pracy wygenerowana'
    except Exception as e:
        errorLabel['text'] = 'Najpierw wygeneruj wyniki'

font = ('Book Antiqua', 15); bg='#ec8c26'
root['bg'] = bg
Label(text='Kodowanie DNA - generator', bg=bg, font=('Book Antiqua', 15, 'bold')).place(x=0, relwidth=1, y=0, relheight=.1)
Label(text='Liczba losowych liter na początku', bg=bg, font=font).place(relx=.1, relwidth=.5, relheight=.1, rely=.1)
frontSpin = Spinbox(from_=0, to=10, font=font, bg='#d6852e'); frontSpin.place(relx=.7, rely=.1, relwidth=.2, relheight=.1)
Label(text='Liczba losowych liter na końcu', bg=bg, font=font).place(relx=.1, relwidth=.5, relheight=.1, rely=.2)
backSpin = Spinbox(from_=0, to=10, font=font, bg='#d6852e'); backSpin.place(relx=.7, rely=.2, relwidth=.2, relheight=.1)
Label(text='Wpisz frazy do zakodowania (max 20 różnych znaków)', bg=bg, font=font).\
place(relx=0, relwidth=1, relheight=.1, rely=.325)
wordsEntry = scrolledtext.ScrolledText(font=("Book Antiqua", 12), bg='#d6852e')
wordsEntry.place(relx=0.1, rely=.425, relwidth=.8, relheight=.375)
errorLabel = Label(text='', bg=bg, font=("Book Antiqua", 12))
errorLabel.place(relx=0, relwidth=1, rely=.8, relheight=.075)
Label(text='Stanisław Piechota wer. pomarańczowa (2021)', bg=bg, font=('Book Antiqua', 7)).place(relx=0.65, relwidth=.35, rely=.87, relheight=.025)
translate = Button(text='Zakoduj hasła', font=font, bg='#888', fg='white', bd=0.5, command=translateEntry)
translate.place(relx=0, relwidth=.5, rely=.9, relheight=.1)
generate = Button(text='Wygeneruj plik .docx', font=font, bg='#888', fg='white', bd=.5, command=generateFiles)
generate.place(relx=.5, relwidth=.5, rely=.9, relheight=.1)

root.bind('<Control-e>', clear)
root.bind('<Control-z>', translateEntry)
root.bind('<Control-g>', generateFiles)
root.mainloop()
